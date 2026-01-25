#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Markdown转Word工具
支持将Markdown文档（含图片）转换为Word文档

使用方法：
1. 将图片放在 docs/images/ 目录下
2. 在Markdown中使用相对路径引用图片：![描述](images/图片名.png)
3. 运行脚本：python scripts/md-to-word.py docs/产品宣传文档.md
"""

import os
import sys
import re
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    import markdown
    from markdown.extensions import tables, fenced_code
except ImportError:
    print("错误：缺少必要的Python库")
    print("请运行以下命令安装：")
    print("  pip install python-docx markdown")
    sys.exit(1)


def find_images_in_markdown(md_content, md_file_path):
    """查找Markdown中的所有图片引用"""
    images = []
    # 匹配 ![alt](path) 格式
    pattern = r'!\[([^\]]*)\]\(([^)]+)\)'
    matches = re.finditer(pattern, md_content)
    
    md_dir = Path(md_file_path).parent
    
    for match in matches:
        alt_text = match.group(1)
        img_path = match.group(2)
        
        # 处理相对路径
        if not os.path.isabs(img_path):
            full_path = md_dir / img_path
        else:
            full_path = Path(img_path)
        
        if full_path.exists():
            images.append({
                'alt': alt_text,
                'path': str(full_path),
                'match': match
            })
        else:
            print(f"警告：图片文件不存在: {full_path}")
    
    return images


def convert_markdown_to_word(md_file_path, output_path=None):
    """将Markdown文件转换为Word文档"""
    
    # 读取Markdown文件
    with open(md_file_path, 'r', encoding='utf-8') as f:
        md_content = f.read()
    
    # 创建Word文档
    doc = Document()
    
    # 设置中文字体
    def set_chinese_font(run):
        run.font.name = '微软雅黑'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    
    # 设置文档默认字体
    style = doc.styles['Normal']
    font = style.font
    font.name = '微软雅黑'
    font._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    font.size = Pt(11)
    
    # 查找所有图片
    images = find_images_in_markdown(md_content, md_file_path)
    
    # 按行处理Markdown内容
    lines = md_content.split('\n')
    i = 0
    
    while i < len(lines):
        line = lines[i]
        
        # 处理标题
        if line.startswith('# '):
            p = doc.add_heading(line[2:].strip(), level=1)
            for run in p.runs:
                set_chinese_font(run)
        elif line.startswith('## '):
            p = doc.add_heading(line[3:].strip(), level=2)
            for run in p.runs:
                set_chinese_font(run)
        elif line.startswith('### '):
            p = doc.add_heading(line[4:].strip(), level=3)
            for run in p.runs:
                set_chinese_font(run)
        elif line.startswith('#### '):
            p = doc.add_heading(line[5:].strip(), level=4)
            for run in p.runs:
                set_chinese_font(run)
        
        # 处理图片
        elif re.match(r'!\[([^\]]*)\]\(([^)]+)\)', line):
            match = re.match(r'!\[([^\]]*)\]\(([^)]+)\)', line)
            alt_text = match.group(1)
            img_path = match.group(2)
            
            md_dir = Path(md_file_path).parent
            if not os.path.isabs(img_path):
                full_path = md_dir / img_path
            else:
                full_path = Path(img_path)
            
            if full_path.exists():
                try:
                    # 添加图片
                    paragraph = doc.add_paragraph()
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = paragraph.add_run()
                    run.add_picture(str(full_path), width=Inches(6))
                    # 添加图片说明
                    if alt_text:
                        caption = doc.add_paragraph(alt_text)
                        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        caption.style = 'Caption'
                        for run in caption.runs:
                            set_chinese_font(run)
                            run.font.size = Pt(9)
                            run.font.color.rgb = RGBColor(128, 128, 128)
                except Exception as e:
                    print(f"警告：无法插入图片 {full_path}: {e}")
                    doc.add_paragraph(f"[图片: {alt_text}]")
            else:
                doc.add_paragraph(f"[图片不存在: {img_path}]")
        
        # 处理表格
        elif line.startswith('|') and '|' in line:
            # 收集表格行
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                table_lines.append(lines[i])
                i += 1
            i -= 1  # 回退一行
            
            if len(table_lines) >= 2:
                # 解析表头
                headers = [cell.strip() for cell in table_lines[0].split('|')[1:-1]]
                # 创建表格
                table = doc.add_table(rows=1, cols=len(headers))
                table.style = 'Light Grid Accent 1'
                
                # 添加表头
                header_cells = table.rows[0].cells
                for j, header in enumerate(headers):
                    header_cells[j].text = header
                    for run in header_cells[j].paragraphs[0].runs:
                        set_chinese_font(run)
                        run.font.bold = True
                
                # 添加数据行（跳过分隔行）
                for row_line in table_lines[2:]:
                    row_cells = [cell.strip() for cell in row_line.split('|')[1:-1]]
                    if len(row_cells) == len(headers):
                        row = table.add_row()
                        for j, cell_text in enumerate(row_cells):
                            row.cells[j].text = cell_text
                            for run in row.cells[j].paragraphs[0].runs:
                                set_chinese_font(run)
        
        # 处理代码块
        elif line.startswith('```'):
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith('```'):
                code_lines.append(lines[i])
                i += 1
            
            if code_lines:
                code_text = '\n'.join(code_lines)
                paragraph = doc.add_paragraph(code_text)
                paragraph.style = 'No Spacing'
                for run in paragraph.runs:
                    set_chinese_font(run)
                    run.font.name = 'Consolas'
                    run.font.size = Pt(9)
        
        # 处理列表
        elif line.strip().startswith('- ') or line.strip().startswith('* '):
            text = line.strip()[2:].strip()
            p = doc.add_paragraph(text, style='List Bullet')
            for run in p.runs:
                set_chinese_font(run)
        elif re.match(r'^\d+\.\s+', line.strip()):
            text = re.sub(r'^\d+\.\s+', '', line.strip())
            p = doc.add_paragraph(text, style='List Number')
            for run in p.runs:
                set_chinese_font(run)
        
        # 处理分隔线
        elif line.strip() == '---':
            doc.add_paragraph('─' * 50)
        
        # 处理普通段落
        elif line.strip() and not line.startswith('#'):
            # 处理截图标记
            if '【截图' in line:
                # 提取截图标记
                screenshot_match = re.search(r'【截图\d+：([^】]+)】', line)
                if screenshot_match:
                    screenshot_desc = screenshot_match.group(1)
                    # 移除截图标记，添加为注释
                    line = re.sub(r'【截图\d+：[^】]+】', '', line)
                    if line.strip():
                        p = doc.add_paragraph(line.strip())
                        for run in p.runs:
                            set_chinese_font(run)
                        # 添加截图占位符
                        placeholder = doc.add_paragraph(f"[截图位置: {screenshot_desc}]")
                        placeholder.style = 'Intense Quote'
                        for run in placeholder.runs:
                            set_chinese_font(run)
                            run.font.size = Pt(10)
                            run.font.color.rgb = RGBColor(0, 112, 192)
                    else:
                        placeholder = doc.add_paragraph(f"[截图位置: {screenshot_desc}]")
                        placeholder.style = 'Intense Quote'
                        for run in placeholder.runs:
                            set_chinese_font(run)
                            run.font.size = Pt(10)
                            run.font.color.rgb = RGBColor(0, 112, 192)
                else:
                    p = doc.add_paragraph(line.strip())
                    for run in p.runs:
                        set_chinese_font(run)
            else:
                p = doc.add_paragraph(line.strip())
                for run in p.runs:
                    set_chinese_font(run)
        
        i += 1
    
    # 保存文档
    if output_path is None:
        output_path = Path(md_file_path).with_suffix('.docx')
    
    doc.save(str(output_path))
    print(f"✅ Word文档已生成: {output_path}")
    print(f"📄 共处理 {len(images)} 张图片")


if __name__ == '__main__':
    if len(sys.argv) < 2:
        print("使用方法: python scripts/md-to-word.py <markdown文件路径> [输出Word文件路径]")
        print("\n示例:")
        print("  python scripts/md-to-word.py docs/产品宣传文档.md")
        print("  python scripts/md-to-word.py docs/产品宣传文档.md output.docx")
        sys.exit(1)
    
    md_file = sys.argv[1]
    output_file = sys.argv[2] if len(sys.argv) > 2 else None
    
    if not os.path.exists(md_file):
        print(f"错误：文件不存在: {md_file}")
        sys.exit(1)
    
    convert_markdown_to_word(md_file, output_file)
